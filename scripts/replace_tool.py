#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用关键词替换工具
支持: docx, xlsx, doc, xls
"""

import os
import sys
import shutil
from pathlib import Path
from typing import Dict, List, Tuple, Optional

try:
    from docx import Document
except ImportError:
    print("错误: 未安装 python-docx")
    print("请运行: uv add python-docx")
    sys.exit(1)

try:
    from openpyxl import load_workbook, Workbook
    from openpyxl.utils import get_column_letter
except ImportError:
    print("错误: 未安装 openpyxl")
    print("请运行: uv add openpyxl")
    sys.exit(1)

# 可选依赖，用于处理旧格式
try:
    import xlrd
except ImportError:
    xlrd = None
    print("提示: 安装 xlrd 可以支持 .xls 文件转换: uv add xlrd")


class KeywordReplacer:
    """通用关键词替换器"""

    def __init__(self, dry_run: bool = False, backup: bool = True, auto_convert: bool = True):
        """
        初始化替换器

        Args:
            dry_run: 预览模式，不实际修改文件
            backup: 是否创建备份文件
            auto_convert: 是否自动转换旧格式文件
        """
        self.dry_run = dry_run
        self.backup = backup
        self.auto_convert = auto_convert
        self.stats = {
            'files_processed': 0,
            'files_skipped': 0,
            'files_converted': 0,
            'total_replacements': 0
        }

    def replace_in_docx(self, file_path: str, replacements: Dict[str, str]) -> int:
        """
        在 docx 文件中替换关键词（支持跨 run 替换）

        Args:
            file_path: 文件路径
            replacements: 替换字典 {旧: 新}

        Returns:
            替换次数
        """
        try:
            doc = Document(file_path)
            count = 0

            # 替换段落中的文本（支持跨 run）
            for paragraph in doc.paragraphs:
                for old, new in replacements.items():
                    if old in paragraph.text:
                        count += self._replace_in_paragraph(paragraph, old, new)

            # 替换表格中的文本（支持跨 run）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for old, new in replacements.items():
                                if old in paragraph.text:
                                    count += self._replace_in_paragraph(paragraph, old, new)

            # 保存文件
            if not self.dry_run and count > 0:
                if self.backup:
                    backup_path = f"{file_path}.bak"
                    shutil.copy2(file_path, backup_path)
                    print(f"  ✓ 备份: {backup_path}")
                doc.save(file_path)

            return count

        except Exception as e:
            print(f"  ✗ 处理 docx 失败: {e}")
            return 0

    def _replace_in_paragraph(self, paragraph, old_text: str, new_text: str) -> int:
        """
        在段落中替换文本（支持跨 run）

        Args:
            paragraph: 段落对象
            old_text: 旧文本
            new_text: 新文本

        Returns:
            替换次数
        """
        # 构建完整文本并查找位置
        full_text = paragraph.text
        if old_text not in full_text:
            return 0

        count = 0
        start_pos = 0
        while True:
            # 查找关键词位置
            pos = full_text.find(old_text, start_pos)
            if pos == -1:
                break

            # 找到关键词在哪些 run 中
            char_count = 0
            start_run_idx = None
            end_run_idx = None
            start_offset = 0
            end_offset = 0

            for run_idx, run in enumerate(paragraph.runs):
                run_len = len(run.text)
                if char_count <= pos < char_count + run_len:
                    start_run_idx = run_idx
                    start_offset = pos - char_count
                if char_count <= pos + len(old_text) <= char_count + run_len:
                    end_run_idx = run_idx
                    end_offset = pos + len(old_text) - char_count
                    break
                char_count += run_len

            if start_run_idx is not None and end_run_idx is not None:
                # 执行替换
                if start_run_idx == end_run_idx:
                    # 关键词在同一个 run 中
                    run_text = paragraph.runs[start_run_idx].text
                    paragraph.runs[start_run_idx].text = run_text[:start_offset] + new_text + run_text[end_offset:]
                else:
                    # 关键词跨越多个 run
                    # 修改第一个 run
                    first_run = paragraph.runs[start_run_idx]
                    first_run.text = first_run.text[:start_offset] + new_text

                    # 清空中间的 run
                    for i in range(start_run_idx + 1, end_run_idx + 1):
                        if i < end_run_idx:
                            paragraph.runs[i].text = ''
                        else:
                            paragraph.runs[i].text = paragraph.runs[i].text[end_offset:]

                count += 1
                start_pos = pos + len(new_text)
            else:
                break

        return count

    def replace_in_xlsx(self, file_path: str, replacements: Dict[str, str]) -> int:
        """
        在 xlsx 文件中替换关键词

        Args:
            file_path: 文件路径
            replacements: 替换字典 {旧: 新}

        Returns:
            替换次数
        """
        try:
            wb = load_workbook(file_path)
            count = 0

            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for old, new in replacements.items():
                                if old in cell.value:
                                    cell.value = cell.value.replace(old, new)
                                    count += 1

            # 保存文件
            if not self.dry_run and count > 0:
                if self.backup:
                    backup_path = f"{file_path}.bak"
                    shutil.copy2(file_path, backup_path)
                    print(f"  ✓ 备份: {backup_path}")
                wb.save(file_path)

            return count

        except Exception as e:
            print(f"  ✗ 处理 xlsx 失败: {e}")
            return 0

    def convert_doc_to_docx(self, doc_path: str) -> Optional[str]:
        """
        将 .doc 转换为 .docx（使用 LibreOffice 保留格式）

        Args:
            doc_path: .doc 文件路径

        Returns:
            转换后的 .docx 文件路径，失败返回 None
        """
        import subprocess

        docx_path = doc_path.replace('.doc', '.docx')
        output_dir = os.path.dirname(os.path.abspath(doc_path))

        # 检查是否安装了 LibreOffice
        try:
            subprocess.run(['which', 'soffice'], capture_output=True, check=True)
        except subprocess.CalledProcessError:
            try:
                subprocess.run(['which', 'libreoffice'], capture_output=True, check=True)
            except subprocess.CalledProcessError:
                print(f"  ✗ 需要安装 LibreOffice 才能转换 .doc 文件")
                print(f"  安装命令: sudo apt-get install libreoffice-writer-nogui")
                return None

        try:
            # 使用 LibreOffice 转换（保留格式）
            cmd = ['soffice', '--headless', '--convert-to', 'docx',
                   doc_path, '--outdir', output_dir]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

            if result.returncode == 0 and os.path.exists(docx_path):
                return docx_path
            else:
                print(f"  ✗ LibreOffice 转换失败: {result.stderr}")
                return None

        except subprocess.TimeoutExpired:
            print(f"  ✗ 转换超时")
            return None
        except Exception as e:
            print(f"  ✗ 转换 .doc 失败: {e}")
            return None

    def convert_xls_to_xlsx(self, xls_path: str) -> Optional[str]:
        """
        将 .xls 转换为 .xlsx

        Args:
            xls_path: .xls 文件路径

        Returns:
            转换后的 .xlsx 文件路径，失败返回 None
        """
        if xlrd is None:
            print(f"  ✗ 需要安装 xlrd 才能转换 .xls 文件")
            print(f"  安装命令: uv add xlrd")
            return None

        xlsx_path = xls_path.replace('.xls', '.xlsx')

        try:
            # 读取 xls
            book_xls = xlrd.open_workbook(xls_path, formatting_info=False, on_demand=True)

            # 创建新的 xlsx
            book_xlsx = Workbook()

            # 删除默认的 sheet
            if 'Sheet' in book_xlsx.sheetnames:
                book_xlsx.remove(book_xlsx['Sheet'])

            # 复制每个 sheet
            for sheet_idx in range(book_xls.nsheets):
                sheet_xls = book_xls.sheet_by_index(sheet_idx)
                sheet_xlsx = book_xlsx.create_sheet(title=sheet_xls.name)

                # 复制数据
                for row in range(sheet_xls.nrows):
                    for col in range(sheet_xls.ncols):
                        cell = sheet_xls.cell(row, col)
                        sheet_xlsx.cell(row=row+1, column=col+1, value=cell.value)

            book_xls.close()
            book_xlsx.save(xlsx_path)
            return xlsx_path

        except Exception as e:
            print(f"  ✗ 转换 .xls 失败: {e}")
            return None

    def replace_in_doc(self, file_path: str, replacements: Dict[str, str]) -> int:
        """
        处理 .doc 文件（自动转换为 docx 后处理）

        Args:
            file_path: 文件路径
            replacements: 替换字典

        Returns:
            替换次数
        """
        if not self.auto_convert:
            print(f"  ⚠ .doc 格式需要 LibreOffice 支持")
            print(f"  建议: 使用 LibreOffice 转换为 docx 后处理")
            print(f"  命令: soffice --headless --convert-to docx '{file_path}'")
            return 0

        print(f"  → 自动转换 .doc → .docx")
        docx_path = self.convert_doc_to_docx(file_path)

        if docx_path:
            print(f"  ✓ 已转换为: {docx_path}")
            self.stats['files_converted'] += 1

            # 转换后进行替换
            return self.replace_in_docx(docx_path, replacements)
        else:
            return 0

    def replace_in_xls(self, file_path: str, replacements: Dict[str, str]) -> int:
        """
        处理 .xls 文件（自动转换为 xlsx 后处理）

        Args:
            file_path: 文件路径
            replacements: 替换字典

        Returns:
            替换次数
        """
        if not self.auto_convert:
            print(f"  ⚠ .xls 格式建议转换为 xlsx 后处理")
            print(f"  命令: soffice --headless --convert-to xlsx '{file_path}'")
            return 0

        print(f"  → 自动转换 .xls → .xlsx")
        xlsx_path = self.convert_xls_to_xlsx(file_path)

        if xlsx_path:
            print(f"  ✓ 已转换为: {xlsx_path}")
            self.stats['files_converted'] += 1

            # 转换后进行替换
            return self.replace_in_xlsx(xlsx_path, replacements)
        else:
            return 0

    def replace_in_file(self, file_path: str, replacements: Dict[str, str]) -> int:
        """
        根据文件扩展名选择相应的替换方法

        Args:
            file_path: 文件路径
            replacements: 替换字典 {旧: 新}

        Returns:
            替换次数
        """
        file_path = os.path.abspath(file_path)

        if not os.path.exists(file_path):
            print(f"✗ 文件不存在: {file_path}")
            self.stats['files_skipped'] += 1
            return 0

        ext = Path(file_path).suffix.lower()

        print(f"\n处理: {file_path}")

        if ext == '.docx':
            count = self.replace_in_docx(file_path, replacements)
        elif ext == '.xlsx':
            count = self.replace_in_xlsx(file_path, replacements)
        elif ext == '.doc':
            count = self.replace_in_doc(file_path, replacements)
        elif ext == '.xls':
            count = self.replace_in_xls(file_path, replacements)
        else:
            print(f"  ✗ 不支持的文件类型: {ext}")
            self.stats['files_skipped'] += 1
            return 0

        if count > 0:
            print(f"  ✓ 替换 {count} 处")
            self.stats['files_processed'] += 1
            self.stats['total_replacements'] += count
        elif count == 0:
            print(f"  - 未找到匹配的关键词")
            self.stats['files_skipped'] += 1

        return count

    def list_directory_files(
        self,
        directory: str,
        pattern: str = "*",
        recursive: bool = False
    ) -> List[Path]:
        """
        列出目录中的所有 Office 文件

        Args:
            directory: 目录路径
            pattern: 文件匹配模式
            recursive: 是否递归处理子目录

        Returns:
            文件路径列表
        """
        path = Path(directory)
        glob_method = path.rglob if recursive else path.glob
        extensions = ['.docx', '.xlsx', '.doc', '.xls']

        files = []
        for file_path in glob_method(pattern):
            if file_path.suffix.lower() in extensions:
                files.append(file_path)

        return files

    def print_file_list(self, files: List[Path], directory: str) -> None:
        """
        打印文件列表

        Args:
            files: 文件路径列表
            directory: 基础目录路径
        """
        print("\n" + "=" * 70)
        print("📋 找到的 Office 文件")
        print("=" * 70)

        if not files:
            print("⚠ 未找到任何 Office 文件")
            return

        # 按目录分组
        from collections import defaultdict
        grouped = defaultdict(list)
        base_path = Path(directory).absolute()

        for file_path in files:
            rel_path = file_path.relative_to(base_path)
            parent = rel_path.parent
            grouped[str(parent)].append(rel_path.name)

        # 打印分组文件
        for dir_name in sorted(grouped.keys()):
            print(f"\n📁 {dir_name if dir_name != '.' else '根目录'}")
            for filename in sorted(grouped[dir_name]):
                print(f"   • {filename}")

        print("\n" + "=" * 70)
        print(f"总计: {len(files)} 个文件")
        print("=" * 70)

    def replace_in_directory(
        self,
        directory: str,
        replacements: Dict[str, str],
        pattern: str = "*",
        recursive: bool = False,
        list_first: bool = True
    ) -> None:
        """
        批量处理目录中的文件

        Args:
            directory: 目录路径
            replacements: 替换字典
            pattern: 文件匹配模式
            recursive: 是否递归处理子目录
            list_first: 是否先列出文件
        """
        # 先列出所有文件
        files = self.list_directory_files(directory, pattern, recursive)

        if list_first:
            self.print_file_list(files, directory)

            if not self.dry_run:
                import time
                print("\n⏳ 即将开始处理文件...")
                time.sleep(0.5)  # 给用户时间查看文件列表

        # 处理文件
        for file_path in files:
            self.replace_in_file(str(file_path), replacements)

    def print_summary(self):
        """打印统计摘要"""
        print("\n" + "=" * 50)
        print("处理摘要")
        print("=" * 50)
        print(f"处理文件数: {self.stats['files_processed']}")
        print(f"跳过文件数: {self.stats['files_skipped']}")
        if self.stats.get('files_converted', 0) > 0:
            print(f"转换文件数: {self.stats['files_converted']}")
        print(f"总替换次数: {self.stats['total_replacements']}")
        if self.dry_run:
            print("\n⚠ 预览模式：未实际修改文件")
        print("=" * 50)


def main():
    """命令行入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description='通用关键词替换工具 - 支持 docx, xlsx, doc, xls',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 单文件替换
  python replace_keywords.py file.docx --replace "旧=新" "旧2=新2"

  # 批量处理
  python replace_keywords.py ./docs --replace "公司=新公司" --recursive

  # 预览模式
  python replace_keywords.py file.xlsx --replace "2024=2025" --dry-run

  # 不创建备份
  python replace_keywords.py file.docx --replace "张三=李四" --no-backup
        """
    )

    parser.add_argument('path', help='文件或目录路径')
    parser.add_argument('--replace', '-r', action='append', dest='replacements',
                       help='替换规则，格式: "旧文本=新文本"', required=True)
    parser.add_argument('--recursive', '-R', action='store_true',
                       help='递归处理子目录')
    parser.add_argument('--dry-run', '-n', action='store_true',
                       help='预览模式，不实际修改文件')
    parser.add_argument('--no-backup', action='store_true',
                       help='不创建备份文件')
    parser.add_argument('--no-convert', action='store_true',
                       help='不自动转换旧格式文件（.doc, .xls）')
    parser.add_argument('--no-list', action='store_true',
                       help='处理前不列出文件（默认会先列出所有文件）')

    args = parser.parse_args()

    # 解析替换规则
    replacements = {}
    for rule in args.replacements:
        if '=' not in rule:
            print(f"✗ 无效的替换规则: {rule}")
            print("  格式应为: 旧文本=新文本")
            sys.exit(1)
        old, new = rule.split('=', 1)
        replacements[old] = new

    print(f"替换规则: {len(replacements)} 条")
    for old, new in replacements.items():
        print(f"  '{old}' → '{new}'")

    # 创建替换器
    replacer = KeywordReplacer(
        dry_run=args.dry_run,
        backup=not args.no_backup,
        auto_convert=not args.no_convert
    )

    # 处理
    path = Path(args.path)
    if path.is_file():
        replacer.replace_in_file(args.path, replacements)
    elif path.is_dir():
        replacer.replace_in_directory(
            args.path,
            replacements,
            recursive=args.recursive,
            list_first=not args.no_list
        )
    else:
        print(f"✗ 路径不存在: {args.path}")
        sys.exit(1)

    replacer.print_summary()


if __name__ == '__main__':
    main()
